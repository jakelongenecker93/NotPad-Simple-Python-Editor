import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import keyword
import re
import os
import sys
import bisect


class CodeEditor(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("NotPad Editor")
        self.geometry("900x700")

        # Indentation handling
        self.indent_width = 4
        self.current_file = None

        # Syntax coloring toggle (on by default)
        self.syntax_enabled = tk.BooleanVar(value=True)

        # Remember last search terms
        self.last_find = ""       # for Ctrl+F / Ctrl+Up/Down (blue)
        self.last_highlight = ""  # for Ctrl+H (yellow)

        # Font / zoom settings
        self.font_family = "Consolas"
        self.base_font_size = 11   # treated as 100%
        self.zoom_percent = 100    # updated by Ctrl+Plus / Ctrl+Minus

        self._create_widgets()
        self._create_menu()
        self._bind_events()

        # Open file passed on command line, if any (for right-click -> Open with)
        if len(sys.argv) > 1:
            path = sys.argv[1]
            if os.path.isfile(path):
                self._open_file(path)

    # ---------------------------
    # Font / zoom helpers
    # ---------------------------
    def _current_font_size(self):
        return max(6, int(round(self.base_font_size * self.zoom_percent / 100.0)))

    def _apply_font(self):
        size = self._current_font_size()
        font_tuple = (self.font_family, size)
        self.text.config(font=font_tuple)
        self.gutter.config(font=font_tuple)
        self._update_gutter()  # also updates status bar

    def _zoom_in(self):
        # Limit zoom to avoid insane sizes
        if self.zoom_percent < 300:
            self.zoom_percent += 10
            self._apply_font()

    def _zoom_out(self):
        if self.zoom_percent > 30:
            self.zoom_percent -= 10
            self._apply_font()

    # ---------------------------
    # Widget creation
    # ---------------------------
    def _create_widgets(self):
        font_tuple = (self.font_family, self._current_font_size())

        self.main_frame = tk.Frame(self)
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # Vertical scrollbar shared by gutter and text
        self.v_scroll = tk.Scrollbar(self.main_frame, orient=tk.VERTICAL)
        self.v_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        # Gutter showing indentation counts
        self.gutter = tk.Text(
            self.main_frame,
            width=4,  # enough for up to 3 digits
            padx=2,
            takefocus=0,
            state=tk.DISABLED,
            wrap="none",
            font=font_tuple,
            background="#f0f0f0",
            foreground="#606060",
            borderwidth=0,
            highlightthickness=0,
        )
        self.gutter.pack(side=tk.LEFT, fill=tk.Y)

        # Main text area
        self.text = tk.Text(
            self.main_frame,
            undo=True,
            wrap="none",
            font=font_tuple,
            borderwidth=0,
        )
        self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Make selected text white so it is readable regardless of syntax scheme
        self.text.tag_configure("sel", foreground="white")

        # Connect scrollbar
        self.text.config(yscrollcommand=self._on_textscroll)
        self.gutter.config(yscrollcommand=self._on_gutterscroll)
        self.v_scroll.config(command=self._on_scrollbar)

        # Basic syntax highlight tags
        self.text.tag_configure("py_keyword", foreground="#0000cc")
        self.text.tag_configure("py_builtin", foreground="#0066aa")
        self.text.tag_configure("py_comment", foreground="#008000")
        self.text.tag_configure("py_string", foreground="#aa5500")
        self.text.tag_configure("py_number", foreground="#990099")

        # Highlight-all tag (Ctrl+H) – bright yellow
        self.text.tag_configure("search_highlight", background="yellow")

        # Find-current-match tag (Ctrl+F / Ctrl+Up/Down) – light blue
        self.text.tag_configure("find_match", background="light blue")

        # Ensure selection tag is on top of other tags so white text wins for selection
        self.text.tag_raise("sel")

        # Right-click context menu for edit operations
        self.context_menu = tk.Menu(self, tearoff=False)
        self.context_menu.add_command(label="Cut", command=self._edit_cut)
        self.context_menu.add_command(label="Copy", command=self._edit_copy)
        self.context_menu.add_command(label="Paste", command=self._edit_paste)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Select All", command=self._edit_select_all)

        # Bind right-click on text to show context menu (Button-3 on Windows/Linux)
        self.text.bind("<Button-3>", self._show_context_menu)

        # Status bar (bottom): line info + zoom %
        self.status_bar = tk.Label(self, anchor="w", text="", relief=tk.SUNKEN, padx=4)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        self._update_gutter()  # also updates status bar

    # ---------------------------
    # Menus
    # ---------------------------
    def _create_menu(self):
        menubar = tk.Menu(self)

        # File menu
        file_menu = tk.Menu(menubar, tearoff=False)
        file_menu.add_command(label="New", accelerator="Ctrl+N", command=self._new_file)
        file_menu.add_command(label="Open...", accelerator="Ctrl+O", command=self._open_dialog)
        file_menu.add_command(label="Save", accelerator="Ctrl+S", command=self._save_file)
        file_menu.add_command(label="Save As...", command=self._save_as)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.destroy)
        menubar.add_cascade(label="File", menu=file_menu)

        # Edit menu (cut/copy/paste/select all)
        edit_menu = tk.Menu(menubar, tearoff=False)
        edit_menu.add_command(label="Cut", accelerator="Ctrl+X", command=self._edit_cut)
        edit_menu.add_command(label="Copy", accelerator="Ctrl+C", command=self._edit_copy)
        edit_menu.add_command(label="Paste", accelerator="Ctrl+V", command=self._edit_paste)
        edit_menu.add_separator()
        edit_menu.add_command(label="Select All", accelerator="Ctrl+A", command=self._edit_select_all)
        menubar.add_cascade(label="Edit", menu=edit_menu)

        # Tools menu
        tools = tk.Menu(menubar, tearoff=False)
        tools.add_checkbutton(
            label="Enable Python syntax coloring",
            variable=self.syntax_enabled,
            command=self._toggle_syntax_coloring,
        )
        tools.add_command(
            label="Color Python syntax now",
            accelerator="F5",
            command=self._color_python_now,
        )
        menubar.add_cascade(label="Tools", menu=tools)

        # Export menu
        export_menu = tk.Menu(menubar, tearoff=False)
        export_menu.add_command(label="Export to Word (.docx)", command=self._export_to_word)
        export_menu.add_command(label="Export to PDF (.pdf)", command=self._export_to_pdf)
        menubar.add_cascade(label="Export", menu=export_menu)

        # Info / help dropdown
        info_menu = tk.Menu(menubar, tearoff=False)
        info_menu.add_command(label="Editor info / shortcuts", command=self._show_info)
        menubar.add_cascade(label="Info", menu=info_menu)

        self.config(menu=menubar)

    def _show_info(self):
        """Show a small help/about dialog with shortcuts and features."""
        info_text = (
            "MiniPy / NotPad Editor – quick reference\n\n"
            "Basic file operations:\n"
            "  Ctrl+N   – New file\n"
            "  Ctrl+O   – Open file\n"
            "  Ctrl+S   – Save file\n\n"
            "Edit operations:\n"
            "  Edit menu / r-clk on text area:\n"
            "  Cut (Ctrl+X)\n"
            "  Copy (Ctrl+C)\n"
            "  Paste (Ctrl+V)\n"
            "  Select All (Ctrl+A)\n\n"
            "Indentation:\n"
            "  Tab / Ctrl+Tab  – Indent / Indent selection\n"
            "  Shift+Tab       – Outdent / Outdent selection\n\n"
            "Zoom / font size:\n"
            "  Ctrl+Plus (+)   – Zoom in. Increase font size by 10%.\n"
            "  Ctrl+Minus (-)  – Zoom out. Decrease font size by 10%.\n"
            "  Status bar (bottom) shows current line and zoom.\n\n"
            "Syntax coloring:\n"
            "  Tools → Python syntax coloring   – Turn color on/off.\n\n"
            "Search & replace:\n"
            "  Ctrl+F    – Find text (light blue).\n"
            "  Ctrl+↕    – Find next / previous occurrence.\n"
            "  Ctrl+H    – Highlight all (bright yellow).\n"
            "  Ctrl+R    – Replace all occurrences of a string.\n"
            "Export:\n"
            "  Export to Word (.docx) – 10pt Word doc.\n"
            "  Export to PDF (.pdf)   – 10pt PDF.\n\n"
            "Other notes:\n"
            "  • Left gutter - Number of leading spaces on line.\n"
        )
        messagebox.showinfo("Editor info / shortcuts", info_text)

    # ---------------------------
    # Event bindings
    # ---------------------------
    def _bind_events(self):
        # Keep gutter updated when text changes or widget resizes
        self.text.bind("<<Modified>>", self._on_modified)
        self.text.bind("<Configure>", lambda e: self._update_gutter())

        # Tab / Shift-Tab for indent / outdent (and Ctrl+Tab variants)
        self.text.bind("<Tab>", self._indent_selection)
        self.text.bind("<Shift-Tab>", self._outdent_selection)
        self.text.bind("<Control-Tab>", self._indent_selection)
        self.text.bind("<Control-Shift-Tab>", self._outdent_selection)

        # Keyboard shortcuts for file ops
        self.bind("<Control-n>", lambda e: (self._new_file(), "break"))
        self.bind("<Control-o>", lambda e: (self._open_dialog(), "break"))
        self.bind("<Control-s>", lambda e: (self._save_file(), "break"))
        self.bind("<F5>",       lambda e: (self._color_python_now(), "break"))

        # Edit shortcuts (explicit, even though Tk has some defaults)
        self.bind("<Control-x>", lambda e: (self._edit_cut(), "break"))
        self.bind("<Control-c>", lambda e: (self._edit_copy(), "break"))
        self.bind("<Control-v>", lambda e: (self._edit_paste(), "break"))
        self.bind("<Control-a>", lambda e: (self._edit_select_all(), "break"))

        # Find / highlight / replace
        self.bind("<Control-f>", self._find_text)   # prompt + first find (light blue)
        self.bind("<Control-h>", self._highlight_all)
        self.bind("<Control-r>", self._replace_all)

        # Arrow-based find navigation (using Ctrl+Up / Ctrl+Down)
        self.bind("<Control-Down>", self._find_next)
        self.bind("<Control-Up>", self._find_prev)

        # Zoom controls
        self.bind("<Control-plus>",       lambda e: (self._zoom_in(), "break"))
        self.bind("<Control-KP_Add>",     lambda e: (self._zoom_in(), "break"))
        self.bind("<Control-=>",          lambda e: (self._zoom_in(), "break"))
        self.bind("<Control-minus>",      lambda e: (self._zoom_out(), "break"))
        self.bind("<Control-KP_Subtract>", lambda e: (self._zoom_out(), "break"))

        # Update status bar on cursor movement
        self.text.bind("<KeyRelease>", lambda e: self._update_status_bar())
        self.text.bind("<ButtonRelease-1>", lambda e: self._update_status_bar())

    # ---------------------------
    # Scrolling sync
    # ---------------------------
    def _on_textscroll(self, *args):
        self.gutter.yview_moveto(args[0])
        self.v_scroll.set(*args)

    def _on_gutterscroll(self, *args):
        self.text.yview_moveto(args[0])
        self.v_scroll.set(*args)

    def _on_scrollbar(self, *args):
        self.text.yview(*args)
        self.gutter.yview(*args)

    # ---------------------------
    # Context menu / Edit actions
    # ---------------------------
    def _show_context_menu(self, event):
        # Put focus on the text widget before showing menu
        self.text.focus_set()
        try:
            self.context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()

    def _edit_cut(self):
        try:
            self.text.event_generate("<<Cut>>")
        except tk.TclError:
            pass

    def _edit_copy(self):
        try:
            self.text.event_generate("<<Copy>>")
        except tk.TclError:
            pass

    def _edit_paste(self):
        try:
            self.text.event_generate("<<Paste>>")
        except tk.TclError:
            pass

    def _edit_select_all(self):
        self.text.tag_add("sel", "1.0", "end-1c")
        self.text.mark_set("insert", "1.0")
        self.text.see("insert")
        self._update_status_bar()

    # ---------------------------
    # Status bar
    # ---------------------------
    def _update_status_bar(self):
        try:
            idx = self.text.index("insert")
        except tk.TclError:
            idx = "1.0"
        line_str, col_str = idx.split(".")
        try:
            total_lines = int(self.text.index("end-1c").split(".")[0])
        except tk.TclError:
            total_lines = 1
        self.status_bar.config(
            text=f"Ln {line_str} / {total_lines}   Zoom: {self.zoom_percent}%"
        )

    # ---------------------------
    # File operations
    # ---------------------------
    def _new_file(self):
        if not self._ok_to_discard():
            return
        self.text.delete("1.0", tk.END)
        self.current_file = None
        self.title("MiniPy Editor")
        self._clear_syntax()
        self._clear_search_highlights()
        self._clear_find_match()
        self.last_find = ""
        self.last_highlight = ""
        self._update_gutter()

    def _open_dialog(self):
        if not self._ok_to_discard():
            return
        path = filedialog.askopenfilename(
            filetypes=[
                ("Python files", "*.py *.pyw"),
                ("Text files", "*.txt"),
                ("All files", "*.*"),
            ]
        )
        if path:
            self._open_file(path)

    def _open_file(self, path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(path, "r", encoding="latin-1") as f:
                content = f.read()

        self.text.delete("1.0", tk.END)
        self.text.insert("1.0", content)
        self.current_file = path
        self.title(f"MiniPy Editor - {os.path.basename(path)}")
        self.text.edit_reset()
        self.text.edit_modified(False)

        self._clear_search_highlights()
        self._clear_find_match()

        # Clear any old syntax colors and re-apply depending on toggle
        if self.syntax_enabled.get():
            self._highlight_python()
        else:
            self._clear_syntax()
        self._update_gutter()

    def _save_file(self):
        if self.current_file is None:
            return self._save_as()
        try:
            content = self.text.get("1.0", tk.END)
            with open(self.current_file, "w", encoding="utf-8") as f:
                f.write(content.rstrip("\n") + "\n")
            self.text.edit_modified(False)
        except Exception as e:
            messagebox.showerror("Save error", str(e))

    def _save_as(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".py",
            filetypes=[
                ("Python files", "*.py *.pyw"),
                ("Text files", "*.txt"),
                ("All files", "*.*"),
            ],
        )
        if not path:
            return
        self.current_file = path
        self._save_file()
        self.title(f"MiniPy Editor - {os.path.basename(path)}")

    def _ok_to_discard(self):
        if self.text.edit_modified():
            response = messagebox.askyesnocancel(
                "Unsaved changes", "Save changes before continuing?"
            )
            if response is None:  # Cancel
                return False
            if response:  # Yes
                self._save_file()
        return True

    # ---------------------------
    # Export to Word / PDF
    # ---------------------------
    def _get_pdf_color_from_tags(self, tags):
        """
        Given the list of Tk tags at a character position, return an (r, g, b)
        triple in the 0–1 range for ReportLab. Defaults to black if no syntax tag.
        """
        # Only care about the syntax tags, ignore search/selection highlights
        if "py_comment" in tags:
            hex_color = "#008000"
        elif "py_string" in tags:
            hex_color = "#aa5500"
        elif "py_keyword" in tags:
            hex_color = "#0000cc"
        elif "py_builtin" in tags:
            hex_color = "#0066aa"
        elif "py_number" in tags:
            hex_color = "#990099"
        else:
            # No syntax tag → plain black
            return (0.0, 0.0, 0.0)

        r = int(hex_color[1:3], 16) / 255.0
        g = int(hex_color[3:5], 16) / 255.0
        b = int(hex_color[5:7], 16) / 255.0
        return (r, g, b)

    def _get_word_rgb_from_tags(self, tags):
        """
        Given the list of Tk tags at a character position, return an (r, g, b)
        triple in the 0–255 range for python-docx. Defaults to black if no syntax tag.
        """
        if "py_comment" in tags:
            hex_color = "#008000"
        elif "py_string" in tags:
            hex_color = "#aa5500"
        elif "py_keyword" in tags:
            hex_color = "#0000cc"
        elif "py_builtin" in tags:
            hex_color = "#0066aa"
        elif "py_number" in tags:
            hex_color = "#990099"
        else:
            return (0, 0, 0)

        r = int(hex_color[1:3], 16)
        g = int(hex_color[3:5], 16)
        b = int(hex_color[5:7], 16)
        return (r, g, b)

    def _export_to_word(self):
        """Export the current buffer to a Word document (.docx) with optional syntax colors."""
        # Late import so the editor still runs even if python-docx isn't installed
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_LINE_SPACING
        except ImportError:
            messagebox.showerror(
                "Export to Word",
                "python-docx is required for Word export.\n"
                "Install it with:\n\n  pip install python-docx"
            )
            return

        # Ask user where to save the file
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
        )
        if not path:
            return

        # If syntax coloring is enabled, refresh it so tags are up to date
        use_colors = False
        if getattr(self, "syntax_enabled", None) is not None and self.syntax_enabled.get():
            self._highlight_python()
            use_colors = True

        try:
            doc = Document()

            # Base style: 10pt Consolas, single spacing, no extra paragraph space
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Consolas"
            font.size = Pt(10)

            pf = style.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Determine how many lines are in the widget
            last_line = int(self.text.index("end-1c").split(".")[0])

            for lineno in range(1, last_line + 1):
                line_start = f"{lineno}.0"
                line_end = f"{lineno}.end"
                line_text = self.text.get(line_start, line_end)

                # Create one paragraph per line
                para = doc.add_paragraph()

                # Enforce no extra spacing on each paragraph, just in case
                pformat = para.paragraph_format
                pformat.space_before = Pt(0)
                pformat.space_after = Pt(0)
                pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE

                # Empty line → just leave a blank paragraph
                if line_text == "":
                    continue

                # If syntax colors are disabled, write the whole line as plain black text
                if not use_colors:
                    run = para.add_run(line_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                    # Word default color is black; no need to set explicitly
                    continue

                # Color-aware export: walk each character and group by color into runs
                prev_color = None  # (r, g, b)
                run_text = ""

                for col in range(len(line_text)):
                    idx = f"{lineno}.{col}"
                    tags = self.text.tag_names(idx)
                    color = self._get_word_rgb_from_tags(tags)

                    if color != prev_color:
                        # Flush the previous run if we have one
                        if run_text:
                            run = para.add_run(run_text)
                            run.font.name = "Consolas"
                            run.font.size = Pt(10)
                            if prev_color is not None:
                                r, g, b = prev_color
                            else:
                                r, g, b = (0, 0, 0)
                            run.font.color.rgb = RGBColor(r, g, b)
                            run_text = ""
                        prev_color = color

                    run_text += line_text[col]

                # Flush any remaining run at end of line
                if run_text:
                    run = para.add_run(run_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                    if prev_color is not None:
                        r, g, b = prev_color
                    else:
                        r, g, b = (0, 0, 0)
                    run.font.color.rgb = RGBColor(r, g, b)

            # Save the document
            doc.save(path)
            messagebox.showinfo("Export to Word", f"Saved Word document:\n{path}")

        except Exception as e:
            messagebox.showerror("Export to Word", f"Error exporting to Word:\n{e}")

    def _export_to_pdf(self):
        # Late import so editor still works if library is missing
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase import pdfmetrics
        except ImportError:
            messagebox.showerror(
                "Export to PDF",
                "reportlab is required for PDF export.\nInstall it with:\n  pip install reportlab",
            )
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Document", "*.pdf")],
        )
        if not path:
            return

        # If syntax coloring is enabled, refresh it so tags are up to date
        use_colors = False
        if self.syntax_enabled.get():
            self._highlight_python()
            use_colors = True

        try:
            c = canvas.Canvas(path, pagesize=letter)
            width, height = letter

            margin = 40
            font_name = "Courier"
            font_size = 10  # default 10pt
            line_height = font_size * 1.2  # simple line spacing

            c.setFont(font_name, font_size)
            y = height - margin

            # How many lines are in the widget?
            last_line = int(self.text.index("end-1c").split(".")[0])

            for lineno in range(1, last_line + 1):
                # New page if we run off the bottom
                if y < margin:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = height - margin

                line_start = f"{lineno}.0"
                line_end = f"{lineno}.end"
                line_text = self.text.get(line_start, line_end)

                # Empty line: just move down
                if line_text == "":
                    y -= line_height
                    continue

                # Plain black text if syntax coloring is off
                if not use_colors:
                    c.setFillColorRGB(0, 0, 0)
                    c.drawString(margin, y, line_text)
                    y -= line_height
                    continue

                # Color-aware rendering
                x = margin
                prev_color = None
                run_text = ""

                for col in range(len(line_text)):
                    idx = f"{lineno}.{col}"
                    tags = self.text.tag_names(idx)
                    color = self._get_pdf_color_from_tags(tags)

                    if color != prev_color:
                        # Flush the previous run
                        if run_text:
                            r, g, b = prev_color
                            c.setFillColorRGB(r, g, b)
                            c.drawString(x, y, run_text)
                            x += pdfmetrics.stringWidth(run_text, font_name, font_size)
                            run_text = ""
                        prev_color = color

                    run_text += line_text[col]

                # Flush any remaining run for the line
                if run_text:
                    r, g, b = prev_color or (0.0, 0.0, 0.0)
                    c.setFillColorRGB(r, g, b)
                    c.drawString(x, y, run_text)

                y -= line_height

            c.save()
            messagebox.showinfo("Export to PDF", f"Saved PDF:\n{path}")
        except Exception as e:
            messagebox.showerror("Export to PDF", str(e))


    # ---------------------------
    # Modified handler / gutter
    # ---------------------------
    def _on_modified(self, event=None):
        if self.text.edit_modified():
            self._update_gutter()
            self._refresh_search_tags()
            self.text.edit_modified(False)

    def _update_gutter(self):
        # Preserve current vertical scroll position so Tab / Shift-Tab don't jump
        yview = self.text.yview()

        # Build indentation counts for each line
        content = self.text.get("1.0", "end-1c")
        lines = content.split("\n")
        gutter_lines = []
        for line in lines:
            # Count leading spaces only; ignore tabs, stop at first non-space/non-tab
            count = 0
            for ch in line:
                if ch == " ":
                    count += 1
                elif ch == "\t":
                    # ignore tabs in the count, but continue
                    continue
                else:
                    break
            if count == 0:
                gutter_lines.append("   ")
            else:
                gutter_lines.append(f"{count:3d}")

        gutter_text = "\n".join(gutter_lines) + ("\n" if gutter_lines else "")

        self.gutter.config(state=tk.NORMAL)
        self.gutter.delete("1.0", tk.END)
        self.gutter.insert("1.0", gutter_text)
        self.gutter.config(state=tk.DISABLED)

        # Restore previous vertical scroll position
        self.text.yview_moveto(yview[0])
        self.gutter.yview_moveto(yview[0])

        # Update status bar
        self._update_status_bar()

    def _refresh_search_tags(self):
        """
        After any text edit, ensure highlights only remain where the text
        still exactly matches their original term.
        """
        # Yellow highlights: search_highlight, based on last_highlight
        if self.last_highlight:
            ranges = list(self.text.tag_ranges("search_highlight"))
            for i in range(0, len(ranges), 2):
                start, end = ranges[i], ranges[i + 1]
                substring = self.text.get(start, end)
                if substring != self.last_highlight:
                    self.text.tag_remove("search_highlight", start, end)
        else:
            # If we no longer have a valid term, remove all yellow
            self._clear_search_highlights()

        # Blue highlight: find_match, based on last_find
        if self.last_find:
            ranges = list(self.text.tag_ranges("find_match"))
            for i in range(0, len(ranges), 2):
                start, end = ranges[i], ranges[i + 1]
                substring = self.text.get(start, end)
                if substring != self.last_find:
                    self.text.tag_remove("find_match", start, end)
        else:
            # No search term: remove any leftover blue
            self._clear_find_match()

    # ---------------------------
    # Indentation helpers (Tab / Shift-Tab)
    # ---------------------------
    def _indent_selection(self, event=None):
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
        except tk.TclError:
            # No selection: insert spaces at cursor
            self.text.insert("insert", " " * self.indent_width)
            self._update_gutter()
            return "break"

        # Preserve the original vertical scroll position
        yview = self.text.yview()

        # Get full lines covering the selection
        start_line = int(start.split(".")[0])
        end_line = int(end.split(".")[0])
        start = f"{start_line}.0"
        end = f"{end_line}.end"

        text_block = self.text.get(start, end)
        lines = text_block.split("\n")
        lines = [(" " * self.indent_width) + line for line in lines]
        new_block = "\n".join(lines)

        self.text.delete(start, end)
        self.text.insert(start, new_block)

        # Keep the (updated) block selected
        new_end = self.text.index(f"{start} + {len(new_block)}c")
        self.text.tag_remove("sel", "1.0", "end")
        self.text.tag_add("sel", start, new_end)

        self._update_gutter()
        # Restore vertical scroll so we don't jump
        self.text.yview_moveto(yview[0])
        self.gutter.yview_moveto(yview[0])

        return "break"  # prevent default Tab behavior

    def _outdent_selection(self, event=None):
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
        except tk.TclError:
            # No selection: remove up to indent_width spaces before cursor on this line
            line_start = self.text.index("insert linestart")
            line_to_cursor = self.text.get(line_start, "insert")
            removable = min(
                self.indent_width, len(line_to_cursor) - len(line_to_cursor.lstrip(" "))
            )
            if removable > 0:
                self.text.delete(f"insert - {removable}c", "insert")
            self._update_gutter()
            return "break"

        # Outdent all selected lines
        text_block = self.text.get(start, end)
        lines = text_block.split("\n")
        new_lines = []
        for line in lines:
            removed = 0
            while removed < self.indent_width and line.startswith(" "):
                line = line[1:]
                removed += 1
            new_lines.append(line)

        new_block = "\n".join(new_lines)
        self.text.delete(start, end)
        self.text.insert(start, new_block)

        # Keep the (updated) block selected
        new_end = self.text.index(f"{start} + {len(new_block)}c")
        self.text.tag_remove("sel", "1.0", "end")
        self.text.tag_add("sel", start, new_end)

        self._update_gutter()
        return "break"

    # ---------------------------
    # Helpers for centering find results
    # ---------------------------
    def _see_center(self, index):
        """Scroll so that 'index' is roughly centered in the text widget."""
        self.text.see(index)
        self.text.update_idletasks()
        info = self.text.dlineinfo(index)
        if info is None:
            return
        x, y, width, height, baseline = info
        if height <= 0:
            return
        text_height = self.text.winfo_height()
        target_y = text_height // 2
        delta_pixels = y - target_y
        lines_to_scroll = int(delta_pixels / float(height))
        if lines_to_scroll != 0:
            self.text.yview_scroll(lines_to_scroll, "units")

    # ---------------------------
    # Find / highlight / replace
    # ---------------------------
    def _clear_search_highlights(self):
        """Remove any existing highlight-all tags (yellow)."""
        self.text.tag_remove("search_highlight", "1.0", tk.END)

    def _clear_find_match(self):
        """Remove the current find match highlight (light blue)."""
        self.text.tag_remove("find_match", "1.0", tk.END)

    def _find_text(self, event=None):
        """Ctrl+F: prompt for text and highlight the first match (light blue).
        Empty text clears blue highlight and resets the current search term.
        """
        pattern = simpledialog.askstring(
            "Find",
            "Find text (leave empty to clear blue highlight):",
            initialvalue=self.last_find or "",
            parent=self,
        )
        if pattern is None:
            # Cancel
            return "break"
        if pattern == "":
            # Clear blue highlight and reset last find term
            self._clear_find_match()
            self.last_find = ""
            return "break"

        self.last_find = pattern
        self._clear_find_match()

        start = self.text.index("insert")
        idx = self.text.search(pattern, start, stopindex=tk.END)
        if not idx:
            # Wrap around from start of file
            idx = self.text.search(pattern, "1.0", stopindex=tk.END)
            if not idx:
                messagebox.showinfo("Find", f'"{pattern}" not found.')
                return "break"

        end_idx = f"{idx}+{len(pattern)}c"
        self.text.tag_add("find_match", idx, end_idx)
        self.text.mark_set("insert", idx)
        self._see_center(idx)

        return "break"

    def _find_next(self, event=None):
        """Ctrl+Down: find next occurrence of last search (light blue)."""
        pattern = self.last_find
        if not pattern:
            # If no previous search, behave like Ctrl+F
            return self._find_text()

        self._clear_find_match()
        start = self.text.index("insert")
        # Move forward at least 1 char so we don't find the same spot again
        start = self.text.index(f"{start}+1c")

        idx = self.text.search(pattern, start, stopindex=tk.END)
        if not idx:
            # Wrap to the top
            idx = self.text.search(pattern, "1.0", stopindex=tk.END)
            if not idx:
                messagebox.showinfo("Find next", f'"{pattern}" not found.')
                return "break"

        end_idx = f"{idx}+{len(pattern)}c"
        self.text.tag_add("find_match", idx, end_idx)
        self.text.mark_set("insert", idx)
        self._see_center(idx)

        return "break"

    def _find_prev(self, event=None):
        """Ctrl+Up: find previous occurrence of last search (light blue)."""
        pattern = self.last_find
        if not pattern:
            # If no previous search, behave like Ctrl+F
            return self._find_text()

        self._clear_find_match()
        # Start just before the cursor so we don't re-find the same match
        start = self.text.index("insert -1c")

        idx = self.text.search(pattern, start, stopindex="1.0", backwards=True)
        if not idx:
            # Wrap from bottom
            idx = self.text.search(pattern, "end-1c", stopindex="1.0", backwards=True)
            if not idx:
                messagebox.showinfo("Find previous", f'"{pattern}" not found.')
                return "break"

        end_idx = f"{idx}+{len(pattern)}c"
        self.text.tag_add("find_match", idx, end_idx)
        self.text.mark_set("insert", idx)
        self._see_center(idx)

        return "break"

    def _highlight_all(self, event=None):
        """Ctrl+H: highlight all matches (bright yellow), or clear them if empty."""
        pattern = simpledialog.askstring(
            "Highlight",
            "Text to highlight (leave empty to clear all yellow highlights):",
            initialvalue=self.last_highlight or self.last_find or "",
            parent=self,
        )
        if pattern is None:
            # Cancel: do nothing
            return "break"
        if pattern == "":
            # Explicitly clear highlights and reset last_highlight
            self._clear_search_highlights()
            self.last_highlight = ""
            return "break"

        self.last_highlight = pattern

        # Clear old highlights and current find-match
        self._clear_search_highlights()
        self._clear_find_match()

        start = "1.0"
        count = 0
        while True:
            idx = self.text.search(pattern, start, stopindex=tk.END)
            if not idx:
                break
            end_idx = f"{idx}+{len(pattern)}c"
            self.text.tag_add("search_highlight", idx, end_idx)
            start = end_idx
            count += 1

        if count == 0:
            messagebox.showinfo("Highlight", f'"{pattern}" not found.')

        return "break"

    def _replace_all(self, event=None):
        """Ctrl+R: replace all occurrences of a string."""
        # Ask for the text to find; default to last find term
        find_text = simpledialog.askstring(
            "Replace All",
            "Find text:",
            initialvalue=self.last_find or "",
            parent=self,
        )
        if find_text is None or find_text == "":
            return "break"

        # Ask for replacement text
        replace_text = simpledialog.askstring(
            "Replace All",
            f"Replace '{find_text}' with:",
            parent=self,
        )
        if replace_text is None:
            return "break"

        content = self.text.get("1.0", tk.END)
        occurrences = content.count(find_text)

        if occurrences == 0:
            messagebox.showinfo("Replace All", f'"{find_text}" not found.')
            return "break"

        # Perform replace
        new_content = content.replace(find_text, replace_text)
        self.text.delete("1.0", tk.END)
        self.text.insert("1.0", new_content)

        # Update last_find to the term we just used; last_highlight unchanged
        self.last_find = find_text
        # Clear highlights; edited text no longer guaranteed to match
        self._clear_search_highlights()
        self._clear_find_match()
        self._update_gutter()
        if self.syntax_enabled.get():
            self._highlight_python()
        else:
            self._clear_syntax()

        messagebox.showinfo(
            "Replace All",
            f"Replaced {occurrences} occurrence{'s' if occurrences != 1 else ''}.",
        )

        return "break"

    # ---------------------------
    # Syntax highlighting
    # ---------------------------
    def _clear_syntax(self):
        for tag in ("py_keyword", "py_builtin", "py_comment", "py_string", "py_number"):
            self.text.tag_remove(tag, "1.0", tk.END)

    def _toggle_syntax_coloring(self):
        """Turn syntax coloring on/off via the Tools menu."""
        if self.syntax_enabled.get():
            self._highlight_python()
        else:
            self._clear_syntax()

    def _color_python_now(self):
        """Apply or clear syntax coloring when F5 (or menu) is used."""
        if self.syntax_enabled.get():
            self._highlight_python()
        else:
            self._clear_syntax()

    def _highlight_python(self):
        content = self.text.get("1.0", "end-1c")
        self._clear_syntax()
        if not content.strip():
            return

        # Precompute line start offsets to convert absolute offsets to Tk indices
        lines = content.split("\n")
        offsets = [0]
        total = 0
        for line in lines:
            total += len(line) + 1  # +1 for newline
            offsets.append(total)

        def offset_to_index(offset):
            # Binary search line index
            line_idx = bisect.bisect_right(offsets, offset) - 1
            col = offset - offsets[line_idx]
            return f"{line_idx + 1}.{col}"

        def mark(pattern, tagname, skip_existing=False):
            for match in pattern.finditer(content):
                start, end = match.span()
                if skip_existing:
                    # Skip if other tags already exist here
                    start_idx = offset_to_index(start)
                    overlapping = self.text.tag_names(start_idx)
                    if any(
                        t in ("py_keyword", "py_builtin", "py_comment", "py_string", "py_number")
                        for t in overlapping
                    ):
                        continue
                self.text.tag_add(tagname, offset_to_index(start), offset_to_index(end))

        # Strings (single or double quotes, naive but works decently)
        string_pattern = re.compile(
            r"('([^'\\]|\\.)*'|\"([^\"\\]|\\.)*\")"
        )
        mark(string_pattern, "py_string")

        # Comments
        comment_pattern = re.compile(r"#.*")
        mark(comment_pattern, "py_comment")

        # Keywords
        kw_pattern = re.compile(r"\b(" + "|".join(keyword.kwlist) + r")\b")
        mark(kw_pattern, "py_keyword", skip_existing=True)

        # Builtins
        builtin_names = dir(__builtins__)
        builtin_pattern = re.compile(
            r"\b(" + "|".join(map(re.escape, builtin_names)) + r")\b"
        )
        mark(builtin_pattern, "py_builtin", skip_existing=True)

        # Numbers
        number_pattern = re.compile(r"\b\d+(\.\d+)?\b")
        mark(number_pattern, "py_number", skip_existing=True)


def main():
    app = CodeEditor()
    app.mainloop()


if __name__ == "__main__":
    main()
